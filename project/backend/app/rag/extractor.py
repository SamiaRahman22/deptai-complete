"""
Document text extraction for PDF, DOCX, and TXT files.
"""

import os
from typing import Tuple
from loguru import logger


def extract_text(file_path: str, file_type: str) -> Tuple[str, int]:
    """
    Extract raw text from a document file.
    Returns (text, page_count).
    """
    if file_type == "pdf":
        return _extract_pdf(file_path)
    elif file_type == "docx":
        return _extract_docx(file_path)
    elif file_type == "txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _extract_pdf(file_path: str) -> Tuple[str, int]:
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            page_count = len(reader.pages)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        text = "\n\n".join(text_parts)
        logger.debug(f"PDF extracted: {page_count} pages, {len(text)} chars")
        return text, page_count
    except ImportError:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")


def _extract_docx(file_path: str) -> Tuple[str, int]:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n\n".join(paragraphs)
        # Estimate pages (roughly 300 words per page)
        word_count = len(text.split())
        page_count = max(1, word_count // 300)
        logger.debug(f"DOCX extracted: ~{page_count} pages, {len(text)} chars")
        return text, page_count
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")


def _extract_txt(file_path: str) -> Tuple[str, int]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    word_count = len(text.split())
    page_count = max(1, word_count // 300)
    return text, page_count


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    import re
    # Remove excessive whitespace
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Remove non-printable characters
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", "", text)
    return text.strip()
